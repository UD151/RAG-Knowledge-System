from langchain_community.document_loaders import TextLoader
from langchain_core.documents import Document as LangchainDoc
from docx import Document


def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    text = []

    for p in doc.paragraphs:
        if p.text.strip():
            text.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text)

    return "\n".join(text)


def load_document(file_path: str):
    if file_path.endswith(".txt"):
        loader = TextLoader(file_path, encoding="utf-8")
        return loader.load()

    text = extract_text_from_docx(file_path)
    return [
        LangchainDoc(
            page_content=text,
            metadata={"source": file_path}
        )
    ]